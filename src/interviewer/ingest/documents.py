from __future__ import annotations

import logging
from pathlib import Path

from ..core.errors import ResumeParseError

logger = logging.getLogger(__name__)

SUPPORTED_SUFFIXES: tuple[str, ...] = (".pdf", ".docx", ".txt", ".md")
_MAX_CHARS = 60_000


def read_document(path: Path) -> str:
    """把简历/JD 文件读成纯文本。只做提取，不做任何语义处理。"""
    if not path.exists():
        raise ResumeParseError(f"文件不存在: {path}")
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _read_pdf(path)
    elif suffix == ".docx":
        text = _read_docx(path)
    elif suffix in (".txt", ".md"):
        text = _read_plain(path)
    else:
        raise ResumeParseError(f"不支持的文件类型: {suffix}")
    cleaned = _normalize(text)
    if not cleaned:
        raise ResumeParseError(f"未能从 {path.name} 中提取到文字，可能是扫描件或纯图片")
    return cleaned[:_MAX_CHARS]


def _read_pdf(path: Path) -> str:
    import pdfplumber

    try:
        chunks: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                chunks.append(page.extract_text() or "")
        return "\n".join(chunks)
    except Exception as exc:
        raise ResumeParseError(f"PDF 解析失败: {exc}") from exc


def _read_docx(path: Path) -> str:
    from docx import Document

    try:
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts)
    except Exception as exc:
        raise ResumeParseError(f"DOCX 解析失败: {exc}") from exc


def _read_plain(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
        except OSError as exc:
            raise ResumeParseError(f"读取失败: {exc}") from exc
    raise ResumeParseError(f"无法识别 {path.name} 的文本编码")


def _normalize(text: str) -> str:
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    result: list[str] = []
    blank = 0
    for line in lines:
        if line.strip():
            blank = 0
            result.append(line)
        else:
            blank += 1
            if blank <= 1:
                result.append("")
    return "\n".join(result).strip()
